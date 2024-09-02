from langchain_openai import OpenAIEmbeddings
from langchain_chroma import Chroma
from langchain_core.documents import Document

from pathlib import Path
from datetime import datetime
import docx
import os
os.environ['OPENAI_API_KEY'] = 'sk-proj-3pSCJqVVn8lRccIMTfXFT3BlbkFJvB0pdZtI3qOHc4EmWKKV'

TODAY = datetime.now().strftime("%y%m%d")

# ベタ打ちされたテキストデータのログフォルダへの保存
def save_raw_data(user_input):
    save_dir = Path("Project log")
    save_dir.mkdir(exist_ok=True)
    save_path = save_dir / f"{TODAY}_raw_text_data.txt" # TODO:user_inputをサマライズしてファイル名にする
    with open(save_path, "w", encoding="utf-8") as f:
            f.write(user_input)
    return save_path


# アップロードされたWordファイルのログフォルダへの保存
def save_uploaded_docx_data(uploaded_file):
    doc = docx.Document(uploaded_file)
    user_input = "\n".join([para.text for para in doc.paragraphs])

    save_dir = Path("Project log")
    save_dir.mkdir(exist_ok=True)

    print(type(uploaded_file))
    print(type(uploaded_file.name))
    print(type(uploaded_file.name))
    save_path = save_dir / f"{TODAY}_{uploaded_file.name}.txt"

    with open(save_path, "wb") as f:
        f.write(user_input)
    return user_input, save_path


# アップロードされたWordファイルのログフォルダへの保存
def save_uploaded_txt_data(uploaded_file):
    user_input = uploaded_file.getvalue().decode("utf-8")

    save_dir = Path("Project log")
    save_dir.mkdir(exist_ok=True)
    save_path = save_dir / f"{TODAY}_{uploaded_file.stem}.txt"

    with open(save_path, "wb") as f:
        f.write(user_input)
    return user_input, save_path


# いずれも1行空きのあるところでチャンク分け
def _data_process(user_input):
    chunks = user_input.split("\n\n")
    return chunks


def _build_vector_database(chunks, save_path):
    # ChromaDBにベクトルデータベースを構築
    chroma_db = Chroma(
        collection_name='projectlog_collection',
        embedding_function=OpenAIEmbeddings(model='text-embedding-ada-002'),
        persist_directory='./.data_projectlog',
    )

    # ベクトルデータベースにドキュメントを追加)
    for idx, chunk in enumerate(chunks):
        print(chunk) # どのチャンクをベクトル化したか確認する用
        document = Document(
            page_content=chunk, # チャンクをベクトル化する
            metadata={
                'file_path': str(save_path),
                'idx': idx}
            )
        chroma_db.add_documents(documents=[document])
    print("ベクトルデータベースの更新が完了しました。")


def build_vector_database(input_folder=Path('Project log')):
    # ベクトル化したいファイルを読み込む
    files = list(input_folder.glob('*.txt'))
    files = [file for file in files if not str(file).startswith("~$")]
    print(files)

    # 1ファイルずつベクトルデータベースを構築
    for file in files:
        df = _data_process(file)
        _build_vector_database(df, file)
        print(f"{file}の処理が完了しました。")


if __name__ == "__main__":
    input_folder=Path('Project log')
    build_vector_database(input_folder)